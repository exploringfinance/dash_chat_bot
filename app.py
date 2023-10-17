###########################
# Author: Tony Trevisan
# Website: www.altanalyticsllc.com
# Date: 2023-10-14
# File: app.py (dash_chat_bot)
# Description: Chat bot that allows a user to upload PDF/TXT files
# Notes: 
###########################

# pip install -r requirements.txt

# Load libraries
import dash
from dash import html, dcc, callback, Output, Input, State
import dash_bootstrap_components as dbc
import boto3
import json
import base64
import io
import os
import PyPDF2
from docx import Document


# Set the AWS credentials file
credentials_file_path = 'assets/credentials'
os.environ['AWS_SHARED_CREDENTIALS_FILE'] = credentials_file_path

# Create app with dbc theme
app = dash.Dash(__name__, external_stylesheets=[dbc.themes.LUMEN]) # SANDSTONE


# Read in the initial/silly prompt in assets folder
silly_prompt = 'assets/silly_prompt.txt'
with open(silly_prompt, 'r') as file:
    silly_prompt = file.read()


# Set the user code based on file so not embedded
user_code = 'assets/user_code.txt'
with open(user_code, 'r') as file:
    user_code = file.read()
code_expected = user_code.strip()

# Function to read in PDF file
def extract_text_from_pdf(contents):
    try:
        content_type, content_string = contents.split(',')
        decoded = base64.b64decode(content_string)

        # Create a PdfFileReader object
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(decoded))

        # Extract text from each page
        text = ""
        for page_number in range(len(pdf_reader.pages)):
            text += pdf_reader.pages[page_number].extract_text()

        return text
    except Exception as e:
        print(f"Error extracting text from PDF: {str(e)}")
        return "Error extracting text"


def parse_contents(contents, filename, date):
    content_type, content_string = contents.split(',')
    decoded = base64.b64decode(content_string)
    try:
        if '.pdf' in filename:
            text_string = extract_text_from_pdf(contents)
        elif '.txt' in filename:
            # Assume that the user uploaded an excel file
            text_string = io.StringIO(decoded.decode('ISO-8859-1')).read()
        elif '.docx' in filename:
          doc = Document(io.BytesIO(decoded))
          text = [paragraph.text for paragraph in doc.paragraphs]
          text_string = '\n'.join(text)
        else:
          text_string = 'Wrong file type uploaded'
          filename = 'Please upload either a PDF, txt, or docx file'
          
    except Exception as e:
        print(e)
        return ['There was an error processing this file.', 'Error with ' + filename]
    return [text_string, filename]


# Create text for the app in Markdown format
md_style = {'width': '80%', 'font-size':'14px'}
md_initial = """
Welcome to the ALT Analytics LLM demo. The application is built in Python Dash and is connected
to the Anthropic Claude 2 model using AWS Bedrock. The app objective is to provide a user interface
for interacting with the LLM without having to go through the AWS console. Furthermore, it gives
users the ability to upload PDF, Word Documents, or TXT files before asking their question. For example,
you can upload several PDF documents and then ask the LLM to summarize them for you or check for 
grammar erros. 

To enable the app, **you must have the correct code**. Reach out to 
[tony@altanalyticsllc.com](mailto:tony@altanalyticsllc.com) to get the code or 
else the app will not work. The user has the ability to enable "Silly Mode" for the app. This 
will usually generate false and funny answers. If the app is in "Silly Mode", you should not take 
seriously any of the output from the model. If you would like to know how you can build a similar 
model on your own, you can follow the instructions laid out in 
[this article](https://www.altanalyticsllc.com/posts/2023-10-15-aws-llm-chatbot/).
"""
md_instructions = dcc.Markdown(children = md_initial, style = md_style, 
                              dangerously_allow_html=True) 
                              
md_upload = """
Use the box below to upload a TXT, DOCX, or PDF file that will be included in the initial prompt. 
**You have to upload the files before making an initial prompt** or it will not work. 
If you have already made an initial request, then you will need to press "Reset" and start over 
in order for your PDF/TXT/DOCX files to be included. To reduce the output size, 
document text will not appear in the response below. 

"""
md_upload_ins = dcc.Markdown(children = md_upload, style = md_style, 
                              dangerously_allow_html=True) 
                              
md_prompt = """
When entering your prompt, be clear about what you are asking and what your response should look like.
If asking the LLM to summarize or review a document. Upload the document(s) and then enter the prompt: 
"Can you summarie the document(s) above" or "Are there any gramatical mistakes in the document(s) above". 
Be sure to reference the "Document(s) above" because the text is parsed and named as "documents". It's 
best to experiment. The model has built in memory, so you continue asking without refreshing the page.
This means, follow-up questions do not have to be as specific. *Enjoy and have fun*!
"""
md_prompt_ins = dcc.Markdown(children = md_prompt, style = md_style, 
                              dangerously_allow_html=True) 


md_valid_code = dcc.Markdown(children = '*Please enter the valid code in the box above*', 
                             style = md_style, dangerously_allow_html=True) 


# Set application layout
app.layout = html.Div([
  # Top navigation bar
  dbc.Nav(
    children = [dbc.NavLink('ALT Analytics LLM Demo',href = 'https://www.altanalyticsllc.com',
                style = {'color': 'white', 'font-size':'20px'})],
                class_name = 'bg-primary', style = {'height': '65px', 'line-height':'45px'}, 
                pills = True),
  dbc.Container([

  # Initial instructions
  html.Div(children = [
    html.Br(),
    html.H4("LLM Powered by AWS and Claude 2 from Anthropic"),
    md_instructions,
    dbc.Row([dbc.Col(dbc.Input(id='code-input', placeholder = 'Enter the code to enable the app', 
                              type = 'text'),width = 4),
             dbc.Col(width=1),
             dbc.Col(dbc.RadioItems(
             options=[
                {"label": "Normal Mode", "value": 1},
                {"label": "Silly Mode", "value": 2}],value=1,
                id="silly-input",), width = 4)
             ]),
    html.Br(),
    
    # Upload document section
    md_upload_ins,
    dbc.Row([dbc.Col(dcc.Upload(
        id='upload-data',
        children=html.Div([
            'Drag and Drop or ',
            html.A('Select Files')
        ]),
        style={
            'width': '100%',
            'height': '60px',
            'lineHeight': '60px',
            'borderWidth': '1px',
            'borderStyle': 'dashed',
            'borderRadius': '5px',
            'textAlign': 'center',
            'margin': '10px',
            'color': '#158cba',  
            'borderColor': '#158cba'  
        },
        # Allow multiple files to be uploaded
        multiple=True
    ), width = 6)
             ]),
    html.Div(id='upload-file-display'),
    dcc.Store(id='upload-file-content'),
    dcc.Store(id='upload-file-name'),
    # Section with the prompt
    html.Hr(),
    md_prompt_ins,
     html.Div(id='question-display'),
     dcc.Store(id = 'prompt-store'),
     dbc.Spinner(html.Div(id='loading-output-chat'), color = "warning"),
     dbc.Row([dbc.Col(dbc.Input(id='question-input', placeholder = 'Enter your prompt here', type = 'text'), width = 8),
              dbc.Col(dbc.Button('Go', id = 'submit-button', color = 'success'), width = 1),
              dbc.Col(dbc.Button('Reset', id = 'reset-button', color = 'danger'), width = 2)]),
   html.Br(),
   html.Br(),
   
   # Bottom nav bar
   html.Div(dbc.NavbarSimple(
    children = [dbc.NavLink('Source Code',href = 'https://github.com/exploringfinance/dash_chat_bot',
                style = {'color': 'white', 'font-size':'12px'})],
                style = {'height': '45px'}, 
                color = 'primary', fixed = 'bottom',
  ))]),
])])


# This is used to replace the raw text output to make it more readable
human = """  

<u>*You*</u>:  
  
"""
assistant = """  

*[Claude2](https://www.anthropic.com/index/claude-2):*  

"""

# Callback to read in text file
@app.callback(Output('upload-file-content', 'data'),
              Output('upload-file-name', 'data'),
              Output('upload-file-display', 'children'),
              Input('code-input','value'),
              Input('upload-data', 'contents'),
              State('upload-data', 'filename'),
              State('upload-data', 'last_modified'),
              State('upload-file-content', 'data'),
              State('upload-file-name', 'data'),
              State('submit-button', 'n_clicks'))
def upload_files(code, list_of_contents, list_of_names, list_of_dates, cur_upload, cur_filnames, n_clicks):
 
     # Check for code
     if code is not None:
        if code.lower() != code_expected:
          print('No code')
          return '', '', md_valid_code
      
     if code is None:
        return '', '', md_valid_code
      
     if n_clicks is not None:
        return '', '', html.P('Please reset before uploading new files.')
    
     if list_of_contents is not None:
       string_build = cur_upload
       name_string = cur_filnames
       parsed_data = [
            parse_contents(c, n, d) for c, n, d in
            zip(list_of_contents, list_of_names, list_of_dates)]
       for item in parsed_data:
          string_build = 'The following text is from a doucment named ' + \
                           item[1] + '\nThe document ends with >>>>>>>>' + \
                           '\n' + string_build + '\n' + item[0] + '>>>>>>>>'
          name_string = name_string + ' / ' + item[1]
       # print(name_string)
       # print(string_build)
       return string_build, name_string, html.P('The following file(s) were uploaded: ' + name_string)
     return '', '', html.P('No Files uploaded')



# Main callback that generates the model output
@app.callback(Output('loading-output-chat', 'children'),
              Output('question-display', 'children'),
              Output('prompt-store', 'data'),
              Output('question-input', 'value'),
              Output('submit-button', 'n_clicks'),
              Output('reset-button', 'n_clicks'),
              Input('submit-button', 'n_clicks'),
              Input('reset-button', 'n_clicks'),
              State('question-input', 'value'),
              State('prompt-store', 'data'),
              State('silly-input', "value"),
              State('upload-file-content', 'data'),
              State('code-input','value'),
              prevent_initial_call=True)
def execute_model(n_clicks, r_clicks, input_value, existing_prompt, silly,
                  upload_file_content, code):
  
  # Print r_clicks and n_clicks
  print(n_clicks)
  print(r_clicks)
  # Check if application code was entered
     # Check for code
  if code is not None:
     if code.lower() != code_expected:
        print('No code')
        return "", md_valid_code, "", "", None, None 
  
  # Reset application
  if r_clicks is not None:
    print('Reset')
    return "", "", "", "", None, None 
  
  # Determine whether to make prompt silly or not by using text fil
  initial_prompt = ''
  if silly == 2:
    initial_prompt = silly_prompt
  print(initial_prompt)
  
  # Connect to bedrock
  session = boto3.Session()
  bedrock = session.client(service_name = 'bedrock-runtime', region_name = 'us-east-1')
  
  if input_value is None: 
    input_value = ''
  
  # Determine which prompt to create
  if n_clicks is None:
    return "", "", "", "", None, None 
  # This builds the initial prompt that will start the chat chain
  elif n_clicks == 1:
    claude_prompt = 'Human: ' + upload_file_content + '\n' +\
    initial_prompt + '\n' + input_value + '\nAssistant:'
  # This adds to the existing prompt
  else:
    claude_prompt = existing_prompt + 'Human:' + input_value + 'Assistant:'
  
  # Commented out to keep from Logging user interactions
  # print('Prompt:')
  # print(claude_prompt)
  body = json.dumps({
    "prompt": claude_prompt,
    "max_tokens_to_sample": 1000,
    "temperature": 0.1,
    "top_p": 0.9,
  })
  # print(claude_prompt)
  # Create API request to the model
  modelId = 'anthropic.claude-v2'
  accept = 'application/json'
  contentType = 'application/json'
  response = bedrock.invoke_model(body = body, modelId = modelId, accept = accept, contentType = contentType)
  response_body = json.loads(response.get('body').read())
  prompt_w_response = claude_prompt + response_body.get('completion')
  # print(prompt_w_response)
  
  # Format output for display in markdown
  output_formatted = prompt_w_response.replace('Human:', human)
  output_formatted = output_formatted.replace('Assistant:', assistant)
  output_formatted = output_formatted.replace(upload_file_content, '')
  md_text = dcc.Markdown(children = output_formatted, style = md_style, 
                         dangerously_allow_html=True) #'<b>bold</b> <u>underline</u>')
                         
  # Return formatted text and raw text
  return "", md_text, prompt_w_response, "", n_clicks, None


if __name__ == '__main__':
    app.run_server(debug=True)

